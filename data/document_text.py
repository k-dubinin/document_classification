"""
Извлечение текста из файлов для классификации и обучения.

Форматы:
  .txt, .md      — обычное чтение;
  .docx         — Word (python-docx);
  .pdf          — текстовый слой (PyMuPDF); для сканов — OCR (Tesseract, при отсутствии - EasyOCR);
  .odt          — OpenDocument Writer (разбор content.xml из ZIP);
  .rtf          — RTF (striprtf);
  .html, .htm   — HTML (извлечение видимого текста, стандартная библиотека).

Старый бинарный .doc (Word 97–2003) не поддерживается — сохраните как .docx.

OCR для PDF:


  pip install pytesseract с установленным tesseract (предпочтительно) или pip install easyocr

"""

from __future__ import annotations

import io
import logging
import os
import zipfile
import xml.etree.ElementTree as ET
from html.parser import HTMLParser

logger = logging.getLogger(__name__)

# --- Параметры по умолчанию (переопределяются через settings/*.yaml, ключи секции ocr.*) ---
TXT_ENCODING = "utf-8"

PDF_TOTAL_TEXT_MIN_FOR_TEXT_ONLY = 120
PDF_PAGE_TEXT_MIN_BEFORE_OCR = 45
PDF_OCR_MATRIX_SCALE = 2.0
EASYOCR_LANGS = ["ru", "en"]
TESSERACT_LANGS = "rus+eng"


def _cfg(key: str, fallback):
    """Значение из YAML/JSON (settings) или запасной константы модуля."""
    try:
        from settings.loader import get_setting

        val = get_setting(key, fallback)
        return fallback if val is None else val
    except Exception:
        return fallback


def extract_text_from_docx(file_path: str) -> str:
    """Текст из документа Word (.docx): абзацы и ячейки таблиц."""
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError("Для .docx установите: pip install python-docx") from e

    document = Document(file_path)
    parts: list[str] = []
    for para in document.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
    return "\n".join(parts)


def _read_file_guess_encoding(file_path: str) -> str:
    """Читает файл как текст, перебирая типичные кодировки."""
    with open(file_path, "rb") as f:
        raw = f.read()
    for enc in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


class _VisibleTextHTMLParser(HTMLParser):
    """Собирает видимый текст из HTML, пропуская script и style."""

    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._skip = False

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag.lower() in ("script", "style"):
            self._skip = True

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in ("script", "style"):
            self._skip = False

    def handle_data(self, data: str) -> None:
        if self._skip:
            return
        t = data.strip()
        if t:
            self._chunks.append(t)


def extract_text_from_html(file_path: str) -> str:
    """Видимый текст из HTML (без скриптов/стилей — только handle_data)."""
    raw = _read_file_guess_encoding(file_path)
    parser = _VisibleTextHTMLParser()
    parser.feed(raw)
    return "\n".join(parser._chunks)


def extract_text_from_rtf(file_path: str) -> str:
    """Текст из RTF."""
    try:
        from striprtf.striprtf import rtf_to_text
    except ImportError as e:
        raise ImportError("Для .rtf установите: pip install striprtf") from e

    raw = _read_file_guess_encoding(file_path)
    return rtf_to_text(raw).strip()


def extract_text_from_odt(file_path: str) -> str:
    """
    Текст из OpenDocument (.odt): чтение content.xml из ZIP без внешних зависимостей.
    """
    chunks: list[str] = []
    with zipfile.ZipFile(file_path, "r") as zf:
        try:
            xml_bytes = zf.read("content.xml")
        except KeyError as e:
            raise ValueError("Файл .odt повреждён: нет content.xml") from e

    root = ET.fromstring(xml_bytes)
    for el in root.iter():
        if el.text and el.text.strip():
            chunks.append(el.text.strip())
        if el.tail and el.tail.strip():
            chunks.append(el.tail.strip())
    return "\n".join(chunks)


def _try_ocr_pdf_page(page) -> str:
    """
    Распознавание текста со страницы PDF как с изображения (скан).
    Сначала пробует Tesseract, при его отсутствии - EasyOCR.
    """
    try:
        import fitz  # PyMuPDF
        from PIL import Image
        import numpy as np
        import io
    except ImportError:
        return ""

    scale = float(_cfg("ocr.PDF_OCR_MATRIX_SCALE", PDF_OCR_MATRIX_SCALE))
    mat = fitz.Matrix(scale, scale)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    img = Image.open(io.BytesIO(pix.tobytes("png")))

    # Попытка использовать Tesseract
    try:
        import pytesseract
        # Получаем текст с изображения с помощью Tesseract
        text = pytesseract.image_to_string(img, lang='rus+eng')
        if text.strip():  # Если Tesseract вернул осмысленный текст
            return text.strip()
    except ImportError:
        logger.info("Tesseract не установлен, используется EasyOCR")
    except Exception as e:
        logger.warning(f"Tesseract ошибка: {e}. Переход к EasyOCR")

    # Если Tesseract не установлен или не сработал, используем EasyOCR
    try:
        import easyocr
    except ImportError:
        logger.warning("EasyOCR не установлен. Установите: pip install easyocr или pytesseract")
        return ""

    # Конвертируем PIL Image в numpy array для EasyOCR
    img_np = np.array(img)

    # Инициализируем reader (с кэшированием для повторных вызовов)
    if not hasattr(_try_ocr_pdf_page, '_reader'):
        langs = _cfg("ocr.EASYOCR_LANGS", EASYOCR_LANGS)
        _try_ocr_pdf_page._reader = easyocr.Reader(langs, gpu=False)  # CPU mode

    try:
        results = _try_ocr_pdf_page._reader.readtext(img_np)
        # Собираем текст из результатов
        text_parts = [result[1] for result in results if result[2] > 0.1]  # confidence > 0.1
        return ' '.join(text_parts)
    except Exception:
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """
    Текст из PDF: сначала текстовый слой; при малом объёме текста — OCR (сканы).

    Для сканов без текстового слоя рекомендуется установить pytesseract,
    при его отсутствии будет использоваться easyocr (см. докстринг модуля).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ImportError("Для .pdf установите: pip install pymupdf") from e

    total_min = int(_cfg("ocr.PDF_TOTAL_TEXT_MIN_FOR_TEXT_ONLY", PDF_TOTAL_TEXT_MIN_FOR_TEXT_ONLY))
    page_min = int(_cfg("ocr.PDF_PAGE_TEXT_MIN_BEFORE_OCR", PDF_PAGE_TEXT_MIN_BEFORE_OCR))

    with fitz.open(file_path) as doc:
        page_texts: list[tuple[object, str]] = []
        for page in doc:
            raw = page.get_text() or ""
            page_texts.append((page, raw))

        total_plain = sum(len(t.strip()) for _, t in page_texts)
        force_ocr_all_pages = total_plain < total_min

        parts: list[str] = []
        ocr_attempts = 0
        ocr_nonempty = 0
        for page, text in page_texts:
            t = text.strip()
            need_ocr = force_ocr_all_pages or len(t) < page_min
            if need_ocr:
                ocr_attempts += 1
                ocr = _try_ocr_pdf_page(page).strip()
                if ocr:
                    ocr_nonempty += 1
                merged = (t + "\n" + ocr).strip() if ocr else t
                parts.append(merged)
            else:
                parts.append(t)

    result = "\n\n".join(p for p in parts if p)
    logger.info(
        "Извлечение PDF: файл=%s страниц=%d символов_текстового_слоя=%d "
        "режим_скан=%s попыток_OCR=%d страниц_с_текстом_OCR=%d итого_символов=%d",
        file_path,
        len(page_texts),
        total_plain,
        force_ocr_all_pages,
        ocr_attempts,
        ocr_nonempty,
        len(result),
    )
    return result


def read_text_from_document(file_path: str, txt_encoding: str = TXT_ENCODING) -> str:
    """
    Читает текст из файла по расширению.

    Parameters
    ----------
    file_path : str
        Путь к поддерживаемому файлу.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"Файл не найден: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext in (".txt", ".md"):
        with open(file_path, encoding=txt_encoding, errors="replace") as f:
            text = f.read()
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".odt":
        text = extract_text_from_odt(file_path)
    elif ext == ".rtf":
        text = extract_text_from_rtf(file_path)
    elif ext in (".html", ".htm"):
        text = extract_text_from_html(file_path)
    else:
        raise ValueError(
            f"Неподдерживаемый формат «{ext}». "
            "Допустимы: .txt, .md, .docx, .pdf, .odt, .rtf, .html, .htm"
        )

    n = len(text or "")
    if ext == ".pdf":
        logger.info("Документ готов: файл=%s формат=%s символов=%d (детали PDF см. выше)", file_path, ext, n)
    else:
        logger.info("Документ прочитан: файл=%s формат=%s символов=%d", file_path, ext, n)
    return text


def supported_document_extensions() -> tuple[str, ...]:
    """Список расширений для обхода папок при обучении."""
    return (
        ".txt",
        ".md",
        ".docx",
        ".pdf",
        ".odt",
        ".rtf",
        ".html",
        ".htm",
    )
